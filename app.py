import streamlit as st
import os
import docx
import shutil
from langchain_community.document_loaders import PyPDFLoader
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_groq import ChatGroq
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from dotenv import load_dotenv

# 1. إعدادات الصفحة
st.set_page_config(page_title="AI Legal Auditor", layout="wide")
load_dotenv()

# تحسين شكل الواجهة
st.markdown("""
    <style>
    .stApp { background-color: #f8f9fa; }
    h1, h2, h3 { color: #1e3d59; font-family: 'Arial'; }
    .stButton>button { width: 100%; border-radius: 8px; background-color: #1e3d59; color: white; height: 3em; font-weight: bold; }
    .stButton>button:hover { background-color: #ffc107; color: #1e3d59; border: 1px solid #1e3d59; }
    .stAlert { border-radius: 10px; }
    </style>
    """, unsafe_allow_html=True)

# دالة قراءة الوورد
def read_docx(file_path):
    doc = docx.Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    return [Document(page_content=text, metadata={"source": "uploaded_docx"})]

@st.cache_resource
def load_models():
    embeddings = HuggingFaceEmbeddings(model_name="./my_model")
    llm = ChatGroq(
        temperature=0, 
        api_key=os.getenv("GROQ_API_KEY"), 
        model_name="llama-3.3-70b-versatile"
    )
    return embeddings, llm

embeddings, llm = load_models()

st.title("⚖️ منصة تدقيق العقود الذكية")
st.write("حلل عقودك بدقة قانونية، اكتشف الثغرات، واضمن حقوقك بميزان القانون المصري.")

uploaded_file = st.file_uploader("ارفع العقد (PDF أو DOCX)", type=["pdf", "docx"])

if uploaded_file:
    file_ext = uploaded_file.name.split('.')[-1].lower()
    temp_path = f"temp_ui.{file_ext}"
    
    with open(temp_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    try:
        # معالجة الملف
        with st.spinner("جاري قراءة وتحليل المستند..."):
            if file_ext == "pdf":
                loader = PyPDFLoader(temp_path)
                docs = loader.load()
            else:
                docs = read_docx(temp_path)

        full_content = " ".join([d.page_content for d in docs])
        
        # --- فحص صحة العقد (Validation) ---
        legal_keywords = ["عقد", "بند", "طرف", "التزام", "اتفاق", "قانون", "صلاحية", "اختصاص", "contract", "agreement"]
        is_legal = any(word in full_content.lower() for word in legal_keywords)

        if not is_legal or len(full_content.strip()) < 150:
            st.error("⚠️ المستند المرفق لا يبدو عقداً قانونياً معتمداً. يرجى رفع ملف يحتوي على بنود قانونية واضحة.")
            st.stop()
        
        # تقسيم النص وإنشاء المستودع الرقمي
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
        splits = text_splitter.split_documents(docs)
        vectorstore = Chroma.from_documents(documents=splits, embedding=embeddings)
        
        # زيادة k لـ 6 لضمان تحليل شامل للمخاطر
        retriever = vectorstore.as_retriever(search_kwargs={"k": 6})

        # دالة المهام الذكية
        def run_legal_task(task_instruction, use_table=False):
            relevant_docs = retriever.invoke(task_instruction)
            context = "\n\n".join([doc.page_content for doc in relevant_docs])
            
            table_info = "برجاء عرض النتائج في شكل جدول Markdown منظم." if use_table else ""

            # الـ Prompt التحليلي المطور
            full_prompt = f"""أنت مستشار قانوني مصري خبير وشديد الذكاء. 
            مهمتك: تحليل النص المرفق والإجابة على سؤال المستخدم بدقة وقوة قانونية.

            قواعد العمل:
            1. استخدم السياق المرفق لتحليل المخاطر، الثغرات، والالتزامات.
            2. في حالة الأسئلة التحليلية (مثل المخاطر)، قم باستنتاج التبعات القانونية بناءً على نصوص العقد والقانون المصري.
            3. إذا وجدت بنداً غامضاً أو معلومة ناقصة (مثل غياب تاريخ الانتهاء أو شروط الفسخ)، وضح ذلك فوراً كخطر محتمل.
            4. الإجابة بالعربية الفصحى فقط، وبأسلوب قانوني رصين.
            5. ممنوع استخدام أي لغة غير العربية في الإجابة.

            السياق المستخرج:
            {context}

            المهمة المطلوبة: {task_instruction}
            {table_info}

            الإجابة التحليلية:"""
            
            with st.spinner("جاري الفحص القانوني..."):
                response = llm.invoke(full_prompt)
                return response.content

        # عرض الواجهة
        st.success("✅ تم التعرف على المستند بنجاح. يمكنك الآن البدء بالتدقيق.")
        
        col1, col2, col3 = st.columns(3)
        with col1:
            if st.button("📝 ملخص العقد"):
                res = run_legal_task("لخص أهم بنود العقد (الأطراف، القيمة، المدة، وطبيعة العمل).", True)
                st.markdown(res)
        with col2:
            if st.button("🚨 كشف المخاطر"):
                res = run_legal_task("استخرج أي ثغرات قانونية أو مخاطر محتملة في هذا العقد بناءً على بنود الفسخ والتعويضات.")
                st.warning(res)
        with col3:
            if st.button("💰 الالتزامات"):
                res = run_legal_task("ما هي الالتزامات المالية، طرق الدفع، والجزاءات المذكورة؟", True)
                st.info(res)

        st.divider()
        st.subheader("💬 اسأل المستشار القانوني")
        user_query = st.text_input("اسأل عن أي بند محدد (مثلاً: ما هو موقف الطرف الثاني في حالة القوة القاهرة؟)")
        if user_query:
            answer = run_legal_task(user_query)
            st.chat_message("assistant").write(answer)

    except Exception as e:
        st.error(f"حدث خطأ أثناء المعالجة: {e}")
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)
else:
    st.info("💡 **نصيحة:** ارفع عقداً واضحاً بصيغة PDF أو Word للحصول على أدق تحليل قانوني.")