from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
import os
import docx
import shutil
from dotenv import load_dotenv
from langchain_community.document_loaders import PyPDFLoader
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_groq import ChatGroq
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

# تحميل المفاتيح من .env
load_dotenv()

app = FastAPI(title="Legal AI Auditor API")

# إعداد الـ CORS عشان الـ Front-end (React/Vue) يقدر يكلم الـ API بدون مشاكل أمنية
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], # في الإنتاج بنحدد دومين الموقع فقط
    allow_methods=["*"],
    allow_headers=["*"],
)

# تحميل النماذج (Embeddings & LLM)
embeddings = HuggingFaceEmbeddings(model_name="./my_model")
llm = ChatGroq(
    temperature=0, 
    api_key=os.getenv("GROQ_API_KEY"), 
    model_name="llama-3.3-70b-versatile"
)

# دالة مساعدة لقراءة ملفات الوورد
def read_docx(file_path):
    doc = docx.Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    return [Document(page_content=text)]

@app.get("/")
def read_root():
    return {"status": "Online", "message": "Legal AI Auditor API is running"}

@app.post("/analyze")
async def analyze_contract(
    file: UploadFile = File(...), 
    query: str = Form(...)
):
    # 1. التأكد من نوع الملف
    file_ext = file.filename.split('.')[-1].lower()
    if file_ext not in ["pdf", "docx"]:
        raise HTTPException(status_code=400, detail="Only PDF and DOCX are supported")

    # 2. حفظ الملف مؤقتاً لمعالجته
    temp_path = f"temp_api.{file_ext}"
    with open(temp_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    try:
        # 3. استخراج النص بناءً على النوع
        if file_ext == "pdf":
            loader = PyPDFLoader(temp_path)
            docs = loader.load()
        else:
            docs = read_docx(temp_path)

        # 4. معالجة النص (Split & Vectorize)
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
        splits = text_splitter.split_documents(docs)
        
        # إنشاء مخزن مؤقت (في الرامات لسرعة الـ API)
        vectorstore = Chroma.from_documents(documents=splits, embedding=embeddings)
        retriever = vectorstore.as_retriever(search_kwargs={"k": 3})

        # 5. البحث عن السياق وسؤال الـ AI
        relevant_docs = retriever.invoke(query)
        context = "\n\n".join([d.page_content for d in relevant_docs])

        full_prompt = f"""أنت مستشار قانوني مصري. بناءً على السياق التالي، أجب على السؤال بدقة وقوة قانونية.
        السياق: {context}
        السؤال: {query}
        الإجابة بالعربية الفصحى:"""

        response = llm.invoke(full_prompt)
        
        return {
            "filename": file.filename,
            "query": query,
            "answer": response.content
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        # مسح الملف المؤقت بعد الانتهاء
        if os.path.exists(temp_path):
            os.remove(temp_path)